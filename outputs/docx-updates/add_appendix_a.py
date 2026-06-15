from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


OUT_DIR = Path(r"D:\CPE491 Project\smart-meter-backend\Project_The_End\outputs\docx-updates")
SOURCE = OUT_DIR / "ระบบติดตามการใช้ไฟฟ้าผ่านมิเตอร์ไฟฟ้า IoT_updated_round1.docx"
TARGET = OUT_DIR / "ระบบติดตามการใช้ไฟฟ้าผ่านมิเตอร์ไฟฟ้า IoT_with_appendix_A.docx"

THAI_FONT = "TH Sarabun New"
CODE_FONT = "Consolas"


def set_run_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)


def shade_paragraph(paragraph, fill="F2F2F2"):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color="B7B7B7", size="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), color)


def add_title(doc, text, size=20):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, THAI_FONT, size, bold=True)
    set_keep_with_next(paragraph)
    return paragraph


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, THAI_FONT, 16, bold=True)
    set_keep_with_next(paragraph)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(36)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run, THAI_FONT, 16)
    return paragraph


def add_file_label(doc, filename):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    label = paragraph.add_run("ชื่อไฟล์: ")
    set_run_font(label, THAI_FONT, 16, bold=True)
    value = paragraph.add_run(filename)
    set_run_font(value, CODE_FONT, 9)
    set_keep_with_next(paragraph)
    return paragraph


def add_code_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, THAI_FONT, 14, bold=True)
    set_keep_with_next(paragraph)
    return paragraph


def add_code(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.right_indent = Pt(12)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.0
    shade_paragraph(paragraph)
    set_paragraph_border(paragraph)
    run = paragraph.add_run(code.strip())
    set_run_font(run, CODE_FONT, 7.5)
    return paragraph


def build_appendix(doc):
    doc.add_page_break()
    add_title(doc, "ภาคผนวก ก", 20)
    add_title(doc, "โค้ดโปรแกรมที่ใช้ในการพัฒนาระบบ", 20)
    add_body(
        doc,
        "ภาคผนวกนี้นำเสนอเฉพาะส่วนสำคัญของโปรแกรมที่ใช้ในการพัฒนาระบบติดตาม"
        "การใช้ไฟฟ้าผ่านมิเตอร์ไฟฟ้า IoT โดยคัดเลือกโค้ดที่แสดงกระบวนการหลักตั้งแต่"
        "การถ่ายภาพและส่งภาพจาก ESP32-CAM การตรวจจับตัวเลขด้วย YOLOv8 "
        "การเลือกเฟรมที่เหมาะสม การบันทึกข้อมูล และการคำนวณบิลค่าไฟฟ้า "
        "ทั้งนี้ไม่ได้แสดงโค้ดทั้งหมดของแต่ละไฟล์ เพื่อให้สามารถอธิบายลำดับการทำงาน"
        "ของระบบได้อย่างกระชับและชัดเจน",
    )

    add_heading(doc, "ก.1 โปรแกรม ESP32-CAM สำหรับถ่ายภาพแบบ Burst")
    add_file_label(doc, "wifi_pic_tune_burst/wifi_pic_tune_burst.ino")
    add_body(
        doc,
        "โปรแกรมส่วนนี้กำหนดรอบการทำงานของ ESP32-CAM ให้ถ่ายภาพหน้าปัดมิเตอร์"
        "จำนวน 10 เฟรม โดยมีระยะห่างเป้าหมายระหว่างเฟรมประมาณ 5 นาที "
        "จากนั้นจึงรวมภาพทั้งหมดและส่งไปยัง Server ในคำขอเดียว รอบการถ่ายภาพครั้งถัดไป"
        "อ้างอิงจากเวลาเริ่มต้นของ Burst เพื่อไม่ให้ระยะเวลาการอัปโหลดทำให้ตารางเวลาเลื่อนออกไป",
    )
    add_code_caption(doc, "โค้ดที่ ก.1 การกำหนดจำนวนเฟรมและรอบการถ่ายภาพ")
    add_code(
        doc,
        """
const unsigned long scheduledIntervalMs = 60UL * 60UL * 1000UL;
const int burstFrameCount = 10;
const unsigned long burstFrameIntervalMs = 5UL * 60UL * 1000UL;

int captureBurstFrames() {
  freeBurstFrames();
  int captured = 0;

  for (int i = 0; i < burstFrameCount; i++) {
    burstFrames[i] = captureFrameCopy(i);
    if (burstFrames[i].ok) captured++;

    if (i < burstFrameCount - 1) {
      waitWithHeartbeat(burstFrameIntervalMs);
    }
  }
  return captured;
}

void runScheduledBurst() {
  unsigned long burstStartMillis = millis();
  lastBurstMillis = burstStartMillis;
  int captured = captureBurstFrames();
  unsigned long burstDurationMs = millis() - burstStartMillis;
  uploadBurstFrames(captured, burstDurationMs);
  freeBurstFrames();
}""",
    )
    add_body(
        doc,
        "ค่าระยะห่างระหว่างเฟรมเป็นค่าคงที่ที่สามารถปรับลดลงระหว่างการทดสอบได้ "
        "แต่การใช้งานตามนโยบายที่กำหนดใช้ระยะห่างประมาณ 5 นาที และส่งค่า "
        "burst_duration_ms เพื่อให้ Backend ประมาณเวลาเริ่มต้นของ Burst ได้",
    )
    add_code_caption(doc, "โค้ดที่ ก.2 การสร้าง Multipart Request และส่งภาพ Burst")
    add_code(
        doc,
        """
bool uploadBurstFrames(int captured, unsigned long burstDurationMs) {
  String boundary = "----ESP32CAMBURST";
  String housePart =
    "--" + boundary + "\\r\\n"
    "Content-Disposition: form-data; name=\\"house_id\\"\\r\\n\\r\\n" +
    String(houseId) + "\\r\\n";
  String durationPart =
    "--" + boundary + "\\r\\n"
    "Content-Disposition: form-data; name=\\"burst_duration_ms\\"\\r\\n\\r\\n" +
    String(burstDurationMs) + "\\r\\n";

  for (int i = 0; i < burstFrameCount; i++) {
    if (!burstFrames[i].ok) continue;
    String imageHead =
      "--" + boundary + "\\r\\n"
      "Content-Disposition: form-data; name=\\"images\\"; "
      "filename=\\"frame_" + String(i + 1) + ".jpg\\"\\r\\n"
      "Content-Type: image/jpeg\\r\\n\\r\\n";
    appendString(payload, offset, imageHead);
    appendBytes(payload, offset, burstFrames[i].data, burstFrames[i].len);
  }

  http.addHeader("Content-Type", "multipart/form-data; boundary=" + boundary);
  int httpCode = http.POST(payload, totalLen);
  return httpCode == 200 || httpCode == -11;
}""",
    )

    add_heading(doc, "ก.2 Backend API สำหรับรับภาพมิเตอร์ไฟฟ้า")
    add_file_label(doc, "routes/upload.js")
    add_body(
        doc,
        "API /api/upload รองรับทั้งภาพเดี่ยวในฟิลด์ image และภาพหลายเฟรมในฟิลด์ images "
        "เมื่อไม่ได้ส่งค่าอ่านแบบ Manual ระบบจะเรียกโปรแกรม Python เพื่อทำนายทุกภาพ "
        "จากนั้นเลือกเฟรมที่เหมาะสมและบันทึกผลที่เลือกเพียงหนึ่งรายการลงในตาราง meter_readings",
    )
    add_code_caption(doc, "โค้ดที่ ก.3 การรับภาพ ทำนายค่า และบันทึกเฟรมที่เลือก")
    add_code(
        doc,
        """
router.post("/", upload.any(), async (req, res) => {
  const imageFiles = (req.files || []).filter(
    (file) => file.fieldname === "image" || file.fieldname === "images"
  );
  const houseId = Number(String(req.body.house_id || "").trim());

  if (!imageFiles.length) {
    return res.status(400).json({ error: "No image uploaded" });
  }
  if (!Number.isInteger(houseId) || houseId <= 0) {
    return res.status(400).json({ error: "house_id is required" });
  }

  const latestReading = await getLatestReading(houseId);
  const predictions = await predictReadingValues(
    imageFiles.map((file) => file.path)
  );
  const frames = imageFiles.map((file, index) => ({
    index,
    filename: file.filename,
    reading_value: normalizePredictedReading(predictions[index]?.reading_value),
    prediction: predictions[index] || null,
  }));

  const selection = chooseBestFrame(frames, latestReading);
  const selected = selection.selected;

  await insertReading({
    houseId,
    readingValue: selected?.reading_value ?? null,
    filename: imageFiles[selected?.index ?? 0].filename,
    captureMode: imageFiles.length > 1 ? "burst" : "single",
    selectedFrame: selected?.index ?? null,
    selectionReason: selection.reason,
    avgConf: selected?.prediction?.avg_conf ?? null,
    framesSummary: summarizeFrames(frames, selected?.index),
  });
});""",
    )

    add_heading(doc, "ก.3 โปรแกรม YOLOv8 สำหรับอ่านค่ามิเตอร์ไฟฟ้า")
    add_file_label(doc, "tools/predict_meter_reading.py")
    add_body(
        doc,
        "โปรแกรม Python โหลดโมเดล YOLOv8 ที่ผ่านการฝึกสอนแล้ว เพื่อตรวจจับตัวเลขบนหน้าปัด"
        "มิเตอร์ จากนั้นเรียงกรอบตรวจจับตามตำแหน่งแกน X จากซ้ายไปขวา ลบกรอบซ้ำที่อยู่ใกล้กัน "
        "และรวมตัวเลขเป็นค่าอ่านมิเตอร์ พร้อมคำนวณค่าความมั่นใจเฉลี่ยของผลตรวจจับ",
    )
    add_code_caption(doc, "โค้ดที่ ก.4 การตรวจจับ เรียงตัวเลข และสร้างค่ามิเตอร์")
    add_code(
        doc,
        """
def summarize_digits(result):
    boxes = result.boxes
    if boxes is None or len(boxes) == 0:
        return "", []

    detections = []
    for xyxy, cls, conf in zip(boxes.xyxy, boxes.cls, boxes.conf):
        x1, y1, x2, y2 = [int(v) for v in xyxy.tolist()]
        detections.append({
            "center_x": (x1 + x2) / 2,
            "digit": int(cls),
            "conf": float(conf),
        })

    detections.sort(key=lambda item: item["center_x"])
    detections = remove_close_duplicates(detections)
    meter_text = "".join(str(item["digit"]) for item in detections)
    return meter_text, detections

model = YOLO(str(MODEL_PATH))
results = model.predict(
    source=[str(path) for path in image_paths],
    conf=0.25,
    iou=0.35,
    agnostic_nms=True,
    verbose=False,
)""",
    )

    add_heading(doc, "ก.4 อัลกอริทึมเลือกเฟรมที่เหมาะสมจาก Burst Upload")
    add_file_label(doc, "routes/upload.js")
    add_body(
        doc,
        "ระบบกรองเฉพาะเฟรมที่อ่านค่าเป็นจำนวนเต็ม และพยายามเลือกค่าที่ไม่ต่ำกว่าค่าอ่านล่าสุด"
        "ของมิเตอร์ จากนั้นจัดกลุ่มเฟรมตามค่าที่อ่านได้ โดยพิจารณาจำนวนเฟรมที่อ่านค่าเดียวกัน "
        "จำนวนตัวเลขที่ตรวจพบครบ 5 หลัก ความใกล้เคียงกับค่ามัธยฐาน และค่าความมั่นใจเฉลี่ย "
        "เมื่อได้กลุ่มที่เหมาะสมแล้วจะเลือกเฟรมที่มีความมั่นใจสูงที่สุดในกลุ่มนั้น",
    )
    add_code_caption(doc, "โค้ดที่ ก.5 หลักการเลือกเฟรมที่น่าเชื่อถือที่สุด")
    add_code(
        doc,
        """
function chooseBestFrame(frames, latestReading) {
  const validFrames = frames.filter((frame) =>
    Number.isInteger(frame.reading_value)
  );
  if (!validFrames.length) {
    return { selected: null, reason: "no_valid_prediction" };
  }

  const monotonicFrames =
    latestReading === null
      ? validFrames
      : validFrames.filter((frame) => frame.reading_value >= latestReading);
  const candidates = monotonicFrames.length ? monotonicFrames : validFrames;
  const values = candidates.map((frame) => frame.reading_value);
  const middleValue = median(values);

  const selectedGroup = buildReadingGroups(candidates).sort((a, b) => {
    if (b.count !== a.count) return b.count - a.count;
    if (b.full_digit_count !== a.full_digit_count) {
      return b.full_digit_count - a.full_digit_count;
    }
    const aDistance = Math.abs(a.reading_value - middleValue);
    const bDistance = Math.abs(b.reading_value - middleValue);
    if (aDistance !== bDistance) return aDistance - bDistance;
    return (b.avg_conf_sum / b.count) - (a.avg_conf_sum / a.count);
  })[0];

  return {
    selected: selectedGroup.best_frame,
    reason: "majority_confidence_median",
  };
}""",
    )

    add_heading(doc, "ก.5 API สำหรับคำนวณและจัดการบิลค่าไฟฟ้า")
    add_file_label(doc, "routes/readings.js")
    add_body(
        doc,
        "การคำนวณบิลใช้ค่าอ่านสะสมของมิเตอร์ในรอบตัดบิลวันที่ 15 ตั้งแต่เวลา 12:00:00 "
        "และก่อนเวลา 13:00:00 โดยใช้ ROW_NUMBER() เลือกค่าอ่านแรกภายในช่วงเวลาของแต่ละเดือน "
        "และใช้ LAG() ดึงค่าอ่านเดือนก่อนหน้า หน่วยไฟฟ้าที่ใช้เท่ากับค่าอ่านเดือนปัจจุบันลบด้วย"
        "ค่าอ่านเดือนก่อนหน้า และจำนวนเงินเท่ากับหน่วยไฟฟ้าที่ใช้คูณด้วยอัตราค่าไฟต่อหน่วย",
    )
    add_code_caption(doc, "โค้ดที่ ก.6 การเลือกค่าอ่านรอบตัดบิลและคำนวณค่าไฟ")
    add_code(
        doc,
        """
WITH cutoff_reading_each_month AS (
  SELECT
    h.id AS house_id,
    DATE_FORMAT(m.reading_time, '%Y-%m-01') AS month_start,
    m.reading_value,
    ROW_NUMBER() OVER (
      PARTITION BY h.id, DATE_FORMAT(m.reading_time, '%Y-%m')
      ORDER BY m.reading_time ASC, m.id ASC
    ) AS row_num
  FROM meter_readings m
  JOIN houses h ON h.id = m.house_id
  WHERE h.is_active = 1
    AND m.reading_value IS NOT NULL
    AND DAY(m.reading_time) = 15
    AND TIME(m.reading_time) >= '12:00:00'
    AND TIME(m.reading_time) < '13:00:00'
),
month_readings AS (
  SELECT
    house_id,
    month_start,
    reading_value AS current_reading,
    LAG(reading_value) OVER (
      PARTITION BY house_id ORDER BY month_start
    ) AS previous_reading
  FROM cutoff_reading_each_month
  WHERE row_num = 1
)
SELECT
  current_reading - previous_reading AS usage_unit,
  ROUND((current_reading - previous_reading) * ?, 2) AS bill_amount
FROM month_readings
WHERE previous_reading IS NOT NULL;""",
    )
    add_code_caption(doc, "โค้ดที่ ก.7 การยกเลิกบิลโดยไม่ลบข้อมูลค่ามิเตอร์")
    add_code(
        doc,
        """
UPDATE electric_bills
SET status = 'cancelled',
    cancelled_at = NOW(),
    cancel_reason = ?
WHERE id = ? AND status <> 'cancelled';""",
    )

    add_heading(doc, "ก.6 โครงสร้างฐานข้อมูลส่วนสำคัญ")
    add_file_label(doc, "smart_meter_db.sql")
    add_body(
        doc,
        "ฐานข้อมูลหลักประกอบด้วยตาราง houses สำหรับข้อมูลบ้าน ตาราง meter_readings "
        "สำหรับค่าอ่านมิเตอร์และข้อมูลการเลือกเฟรม และตาราง electric_bills สำหรับเก็บข้อมูล"
        "บิลที่สร้างแล้วในรูปแบบ Snapshot เพื่อให้สามารถตรวจสอบประวัติและยกเลิกบิลได้โดยไม่กระทบ"
        "ข้อมูลค่ามิเตอร์ต้นทาง",
    )
    add_code_caption(doc, "โค้ดที่ ก.8 โครงสร้างตารางค่ามิเตอร์และประวัติบิล")
    add_code(
        doc,
        """
CREATE TABLE meter_readings (
  id INT NOT NULL AUTO_INCREMENT,
  house_id INT NOT NULL,
  reading_value FLOAT DEFAULT NULL,
  image_filename VARCHAR(255) DEFAULT NULL,
  reading_time TIMESTAMP NULL DEFAULT CURRENT_TIMESTAMP,
  capture_mode VARCHAR(20) DEFAULT 'single',
  selected_frame INT DEFAULT NULL,
  selection_reason VARCHAR(80) DEFAULT NULL,
  avg_conf FLOAT DEFAULT NULL,
  frames_summary LONGTEXT DEFAULT NULL,
  PRIMARY KEY (id),
  FOREIGN KEY (house_id) REFERENCES houses(id)
);

CREATE TABLE electric_bills (
  id INT NOT NULL AUTO_INCREMENT,
  bill_no VARCHAR(80) NOT NULL,
  house_id INT NOT NULL,
  start_month CHAR(7) NOT NULL,
  end_month CHAR(7) NOT NULL,
  start_reading FLOAT NOT NULL,
  end_reading FLOAT NOT NULL,
  usage_unit FLOAT NOT NULL,
  unit_rate DECIMAL(12,4) NOT NULL,
  total_amount DECIMAL(12,2) NOT NULL,
  issue_date DATETIME NOT NULL,
  due_date DATETIME NOT NULL,
  status VARCHAR(20) NOT NULL DEFAULT 'active',
  cancelled_at DATETIME DEFAULT NULL,
  cancel_reason VARCHAR(255) DEFAULT NULL,
  PRIMARY KEY (id),
  FOREIGN KEY (house_id) REFERENCES houses(id)
);""",
    )


def main():
    if not SOURCE.exists():
        raise SystemExit(f"Source document not found: {SOURCE}")

    copy2(SOURCE, TARGET)
    doc = Document(TARGET)
    build_appendix(doc)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
