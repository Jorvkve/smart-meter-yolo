from pathlib import Path

from docx import Document


DOCX = Path(
    r"D:\CPE491 Project\smart-meter-backend\Project_The_End"
    r"\outputs\docx-updates\ระบบติดตามการใช้ไฟฟ้าผ่านมิเตอร์ไฟฟ้า IoT_with_appendix_A.docx"
)
SOURCE = DOCX.with_name("ระบบติดตามการใช้ไฟฟ้าผ่านมิเตอร์ไฟฟ้า IoT_updated_round1.docx")

EXPECTED_TEXT = [
    "ภาคผนวก ก",
    "โค้ดโปรแกรมที่ใช้ในการพัฒนาระบบ",
    "ก.1 โปรแกรม ESP32-CAM สำหรับถ่ายภาพแบบ Burst",
    "ก.2 Backend API สำหรับรับภาพมิเตอร์ไฟฟ้า",
    "ก.3 โปรแกรม YOLOv8 สำหรับอ่านค่ามิเตอร์ไฟฟ้า",
    "ก.4 อัลกอริทึมเลือกเฟรมที่เหมาะสมจาก Burst Upload",
    "ก.5 API สำหรับคำนวณและจัดการบิลค่าไฟฟ้า",
    "ก.6 โครงสร้างฐานข้อมูลส่วนสำคัญ",
]


def main():
    doc = Document(DOCX)
    source_doc = Document(SOURCE)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    source_paragraphs = [paragraph.text for paragraph in source_doc.paragraphs]
    full_text = "\n".join(paragraphs)

    missing = [text for text in EXPECTED_TEXT if text not in full_text]
    captions = [text for text in paragraphs if text.startswith("โค้ดที่ ก.")]
    code_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.runs and paragraph.runs[0].font.name == "Consolas"
    ]
    longest_code_line = max(
        (
            len(line)
            for paragraph in code_paragraphs
            for line in paragraph.text.splitlines()
        ),
        default=0,
    )

    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"source_paragraphs={len(source_doc.paragraphs)}")
    print(f"sections={len(doc.sections)}")
    print(f"appendix_code_captions={len(captions)}")
    print(f"code_paragraphs={len(code_paragraphs)}")
    print(f"longest_code_line={longest_code_line}")
    print(f"missing_expected_text={missing}")

    if missing:
        raise SystemExit("Appendix headings are incomplete")
    if paragraphs[: len(source_paragraphs)] != source_paragraphs:
        raise SystemExit("Original document paragraphs changed unexpectedly")
    if len(captions) != 8:
        raise SystemExit("Expected exactly 8 code captions")
    if len(code_paragraphs) != 8:
        raise SystemExit("Expected exactly 8 code blocks to use Consolas")
    if longest_code_line > 100:
        raise SystemExit("A code line is too long for the document layout")


if __name__ == "__main__":
    main()
